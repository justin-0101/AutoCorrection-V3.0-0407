#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
文档处理模块
用于处理各种格式的文档文件，支持 txt, docx, pdf 和图片格式
"""

import os
import logging
import tempfile
import traceback
from werkzeug.utils import secure_filename

# 获取 logger
logger = logging.getLogger(__name__)

# 支持的文件类型
ALLOWED_EXTENSIONS = {
    'txt': '文本文件',
    'docx': 'Word文档',
    'doc': 'Word文档',
    'pdf': 'PDF文档',
    'jpg': '图片',
    'jpeg': '图片',
    'png': '图片',
    'gif': '图片'
}

def allowed_file(filename):
    """
    检查文件是否为允许的类型
    
    Args:
        filename: 文件名
        
    Returns:
        bool: 文件类型是否允许
    """
    logger.info(f"检查文件格式是否允许: {filename}")
    if '.' not in filename:
        logger.warning(f"文件名 {filename} 没有扩展名")
        return False
    
    ext = filename.rsplit('.', 1)[1].lower()
    logger.info(f"文件扩展名: {ext}, 允许的扩展名: {list(ALLOWED_EXTENSIONS.keys())}")
    return ext in ALLOWED_EXTENSIONS.keys()


def extract_text_from_docx(file_path):
    """
    从.docx 文件中提取文本
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 提取的文本
    """
    try:
        logger.info(f"处理.docx文件: {file_path}")
        
        # 使用 python-docx 库
        try:
            import docx
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
            logger.info(f"成功从.docx文件提取文本，长度: {len(text)}")
            return text
        except ImportError:
            logger.warning("python-docx库未安装，使用替代方法")
            import zipfile
            from xml.etree.ElementTree import XML
            
            WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            PARA = WORD_NAMESPACE + 'p'
            TEXT = WORD_NAMESPACE + 't'
            
            with zipfile.ZipFile(file_path) as docx_zip:
                content = docx_zip.read('word/document.xml')
                tree = XML(content)
                
                paragraphs = []
                for paragraph in tree.iter(PARA):
                    texts = [node.text for node in paragraph.iter(TEXT) if node.text]
                    if texts:
                        paragraphs.append(''.join(texts))
                
                text = "\n".join(paragraphs)
                logger.info(f"成功使用zipfile从.docx文件提取文本，长度: {len(text)}")
                return text
        
    except Exception as e:
        logger.error(f"处理.docx文件时出错: {str(e)}\n{traceback.format_exc()}")
        raise Exception(f"无法读取.docx文件: {str(e)}")


def extract_text_from_txt(file_path):
    """
    从文本文件中提取文本
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 提取的文本
    """
    try:
        logger.info(f"处理.txt文件: {file_path}")
        
        # 尝试不同的编码方式
        encodings = ["utf-8", "gbk", "gb2312", "gb18030", "big5", "latin-1"]
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
                logger.info(f"成功使用{encoding}编码读取文本文件，长度: {len(text)}")
                return text
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"使用{encoding}编码读取失败: {str(e)}")
        
        # 如果所有编码都失败，尝试二进制方式读取
        with open(file_path, 'rb') as f:
            binary_data = f.read()
        
        text = binary_data.decode('utf-8', errors='replace')
        logger.info(f"成功使用二进制方式读取文本文件，长度: {len(text)}")
        return text
        
    except Exception as e:
        logger.error(f"处理.txt文件时出错: {str(e)}\n{traceback.format_exc()}")
        raise Exception(f"无法读取文本文件: {str(e)}")


def extract_text_from_pdf(file_path):
    """
    从PDF文件中提取文本
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 提取的文本
    """
    try:
        logger.info(f"处理.pdf文件: {file_path}")
        
        # 尝试使用阿里千问API进行PDF文本提取
        try:
            from app.core.ai.aliyun_qianwen_client import AliyunQianwenClient
            
            logger.info("使用阿里千问API处理PDF")
            qianwen_client = AliyunQianwenClient()
            text = qianwen_client.extract_text_from_pdf(file_path)
            
            logger.info(f"成功使用阿里千问API从PDF提取文本，长度: {len(text)}")
            return text
        except ImportError as ie:
            logger.warning(f"阿里千问API客户端未找到: {str(ie)}，尝试使用备用方法")
        except Exception as api_e:
            logger.warning(f"阿里千问API调用失败: {str(api_e)}，尝试使用备用方法")
        
        # 备用方法1：PyPDF2
        try:
            import PyPDF2
            
            logger.info("使用PyPDF2处理PDF")
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
                
                logger.info(f"成功从PDF文件提取文本，长度: {len(text)}")
                return text
                
        except ImportError:
            logger.warning("PyPDF2库未安装，尝试使用替代方法")
        except Exception as pdf_e:
            logger.warning(f"PyPDF2处理失败: {str(pdf_e)}，尝试使用替代方法")
            
        # 备用方法2：pdfminer
        try:
            from pdfminer.high_level import extract_text as pdf_extract_text
            
            logger.info("使用pdfminer处理PDF")
            text = pdf_extract_text(file_path)
            logger.info(f"成功使用pdfminer从PDF文件提取文本，长度: {len(text)}")
            return text
        except ImportError:
            logger.warning("pdfminer库未安装")
            raise Exception("未安装PDF处理库，请安装PyPDF2或pdfminer.six")
        except Exception as miner_e:
            logger.warning(f"pdfminer处理失败: {str(miner_e)}")
            raise Exception(f"所有PDF处理方法均失败: {str(miner_e)}")
                
    except Exception as e:
        logger.error(f"处理PDF文件时出错: {str(e)}\n{traceback.format_exc()}")
        raise Exception(f"无法读取PDF文件: {str(e)}")


def extract_text_from_image(file_path):
    """
    从图片中提取文本（OCR）
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 提取的文本
    """
    try:
        logger.info(f"处理图片文件: {file_path}")
        
        # 尝试使用阿里千问API进行OCR
        try:
            from app.core.ai.aliyun_qianwen_client import AliyunQianwenClient
            
            logger.info("使用阿里千问API进行OCR处理")
            qianwen_client = AliyunQianwenClient()
            text = qianwen_client.extract_text_from_image(file_path)
            
            logger.info(f"成功使用阿里千问API从图片提取文本，长度: {len(text)}")
            return text
        except ImportError as ie:
            logger.warning(f"阿里千问API客户端未找到: {str(ie)}，尝试使用备用方法")
        except Exception as api_e:
            logger.warning(f"阿里千问API调用失败: {str(api_e)}，尝试使用备用方法")
            
        # 备用方法：使用pytesseract
        try:
            import pytesseract
            from PIL import Image
            
            logger.info("使用pytesseract进行OCR处理")
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            
            logger.info(f"成功从图片提取文本，长度: {len(text)}")
            return text
            
        except ImportError:
            logger.warning("OCR库未安装")
            raise Exception("未安装OCR处理库，请安装pytesseract和Pillow")
            
    except Exception as e:
        logger.error(f"处理图片文件时出错: {str(e)}\n{traceback.format_exc()}")
        raise Exception(f"无法处理图片文件: {str(e)}")


def extract_text_from_doc(file_path):
    """
    从.doc 文件中提取文本
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 提取的文本
    """
    try:
        logger.info(f"处理.doc文件: {file_path}")
        
        # 尝试直接使用docx2txt
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            logger.info(f"成功使用docx2txt从.doc文件提取文本，长度: {len(text)}")
            return text
        except ImportError:
            logger.warning("docx2txt库未安装，尝试其他方法")
        except Exception as e:
            logger.warning(f"docx2txt处理失败: {str(e)}，尝试其他方法")
        
        # 如果docx2txt失败，尝试使用python-docx（可能对某些.doc文件有效）
        try:
            import docx
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
            logger.info(f"成功使用python-docx从.doc文件提取文本，长度: {len(text)}")
            return text
        except ImportError:
            logger.warning("python-docx库未安装，尝试其他方法")
        except Exception as e:
            logger.warning(f"python-docx处理失败: {str(e)}，尝试其他方法")
        
        # 最后尝试简单地将其视为文本文件
        try:
            text = extract_text_from_txt(file_path)
            logger.info(f"成功使用简单文本方式读取.doc文件，长度: {len(text)}")
            return text
        except Exception as e:
            logger.warning(f"文本方式读取失败: {str(e)}")
        
        # 如果所有方法都失败，抛出异常
        raise Exception("无法使用任何可用方法读取.doc文件")
        
    except Exception as e:
        logger.error(f"处理.doc文件时出错: {str(e)}\n{traceback.format_exc()}")
        raise Exception(f"无法读取.doc文件: {str(e)}")


def process_document(file):
    """
    处理上传的文档文件，提取文本内容
    
    Args:
        file: 上传的文件对象
        
    Returns:
        tuple: (文本内容, 文件标题)
    """
    temp_file = None
    try:
        # 基本验证
        if not file or file.filename == '':
            raise Exception("未选择文件")
            
        # 提取原始文件信息 - 在secure_filename之前
        original_filename = file.filename
        orig_name, orig_ext = os.path.splitext(original_filename)
        orig_ext = orig_ext.lower()  # 转小写但保留点号
        
        # 记录原始文件信息，以备后续使用
        logger.info(f"原始文件名: {original_filename}, 名称: {orig_name}, 扩展名: {orig_ext}")
        
        # 进行安全文件名处理 - 此时可能丢失中文字符
        safe_filename = secure_filename(file.filename)
        logger.info(f"安全文件名: {safe_filename}")
        
        # 检查扩展名是否在允许列表中 - 使用原始扩展名判断
        clean_ext = orig_ext.lstrip('.')  # 去除前导点号
        if not clean_ext or clean_ext not in ALLOWED_EXTENSIONS:
            allowed_exts = ', '.join(ALLOWED_EXTENSIONS.keys())
            logger.error(f"不支持的文件格式: {clean_ext}，支持的格式: {allowed_exts}")
            raise Exception(f"不支持的文件格式。支持的格式：{allowed_exts}")
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(delete=False) as temp:
            temp_file = temp.name
            file.save(temp_file)
        
        # 获取文件标题 - 使用原始名称，确保中文不丢失
        title = orig_name if orig_name else "未命名"
        
        # 根据扩展名选择处理方法
        if clean_ext == 'docx':
            content = extract_text_from_docx(temp_file)
        elif clean_ext == 'doc':
            content = extract_text_from_doc(temp_file)
        elif clean_ext == 'txt':
            content = extract_text_from_txt(temp_file)
        elif clean_ext == 'pdf':
            content = extract_text_from_pdf(temp_file)
        elif clean_ext in ['jpg', 'jpeg', 'png', 'gif']:
            content = extract_text_from_image(temp_file)
        else:
            # 不应该到达这里，因为我们已经检查了扩展名
            raise Exception(f"不支持的文件格式: {clean_ext}")
        
        return content, title
    
    except Exception as e:
        logger.error(f"处理文档时出错: {str(e)}\n{traceback.format_exc()}")
        raise Exception(f"处理文档时出错: {str(e)}")
    
    finally:
        # 清理临时文件
        if temp_file and os.path.exists(temp_file):
            try:
                os.unlink(temp_file)
            except Exception as e:
                logger.error(f"清理临时文件失败: {str(e)}")