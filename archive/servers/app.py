import os
import json
import logging
import traceback
import asyncio
import sqlite3
from sqlite3 import Error as SQLiteError, OperationalError, IntegrityError
from flask import Flask, render_template, redirect, url_for, session, request, flash, g, jsonify, current_app, send_from_directory
from datetime import datetime, timedelta, date
import dotenv
import time
from werkzeug.security import generate_password_hash, check_password_hash
from flask_cors import CORS
from werkzeug.utils import secure_filename
import uuid
import re
from functools import wraps

# 加载环境变量
dotenv.load_dotenv()

# 标记AI服务初始化状态
ai_initialized = False
ai_available = False

# 禁用Werkzeug默认日志记录器以避免双重日志
logging.getLogger('werkzeug').disabled = True

# 从app_instance导入Flask应用实例，确保使用单例模式
from app_instance import app

# 使用环境变量更新app配置
app.secret_key = os.environ.get('SECRET_KEY', 'your_secret_key_here')
app.permanent_session_lifetime = timedelta(days=30)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 限制上传大小为16MB

# 静态文件和模板配置
app.static_folder = 'static'
app.static_url_path = '/static'
app.template_folder = os.path.abspath(os.path.join(os.path.dirname(__file__), 'templates'))

# 确保应用使用的目录存在
if not os.path.exists('instance'):
    os.makedirs('instance')

# 数据库路径
DATABASE = 'instance/essay_correction.db'

# 配置日志
if not os.path.exists('logs'):
    os.makedirs('logs')

# 创建日志格式化器
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')

# 清除根日志记录器的所有处理器
root_logger = logging.getLogger()
for handler in root_logger.handlers[:]:
    root_logger.removeHandler(handler)

# 创建文件处理器
file_handler = logging.FileHandler('logs/app.log', encoding='utf-8')
file_handler.setFormatter(formatter)
file_handler.setLevel(logging.INFO)

# 创建控制台处理器
console_handler = logging.StreamHandler()
console_handler.setFormatter(formatter)
console_handler.setLevel(logging.INFO)

# 配置根日志记录器
root_logger.setLevel(logging.INFO)
root_logger.addHandler(file_handler)
root_logger.addHandler(console_handler)

# 配置 Werkzeug 日志
werkzeug_logger = logging.getLogger('werkzeug')
werkzeug_logger.handlers = []  # 清除现有处理器
werkzeug_logger.addHandler(logging.NullHandler())  # 添加空处理器
werkzeug_logger.propagate = False  # 阻止日志传播到根日志记录器

# 获取应用的日志记录器
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
logger.propagate = False  # 阻止日志传播到根日志记录器

# 清除应用日志记录器的所有处理器
for handler in logger.handlers[:]:
    logger.removeHandler(handler)

# 为应用日志记录器添加处理器
try:
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)
except Exception as e:
    print(f"添加日志处理器时出错: {e}")

# 打印启动信息
logger.info("\n" + "="*50)
logger.info("AI评分系统启动")
logger.info("日志配置完成 - 同时输出到控制台和文件")
logger.info("="*50)

# 添加全局数据库错误处理
@app.errorhandler(OperationalError)
def handle_db_operational_error(e):
    logger.error(f"数据库操作错误: {str(e)}")
    logger.error(traceback.format_exc())
    return jsonify({"error": "数据库操作失败，请联系管理员"}), 500

@app.errorhandler(IntegrityError)
def handle_db_integrity_error(e):
    logger.error(f"数据完整性错误: {str(e)}")
    return jsonify({"error": "数据冲突，请检查输入值"}), 400

@app.errorhandler(SQLiteError)
def handle_db_error(e):
    logger.error(f"SQLite错误: {str(e)}")
    logger.error(traceback.format_exc())
    return jsonify({"error": "数据库错误，请稍后重试"}), 500

@app.errorhandler(Exception)
def handle_unexpected_error(e):
    logger.error(f"未预期的错误: {str(e)}")
    logger.error(traceback.format_exc())
    return jsonify({"error": "服务器内部错误，请稍后重试"}), 500

# 初始化AI评分服务
def initialize_ai_service():
    global ai_initialized, ai_available
    if not ai_initialized:
        try:
            from llm_func import model, get_model_response, handler_spell_error, handler_content_analysis, handler_express_analysis, handler_writing_analysis, handler_summary
            ai_available = True
            print("AI评分服务初始化成功")
            ai_initialized = True
        except Exception as e:
            print(f"AI评分服务初始化失败: {e}")
            ai_available = False
            ai_initialized = True

# 初始化AI服务
initialize_ai_service()

# 统一的AI评分函数
def unified_ai_scoring(title, content):
    """
    统一的AI评分函数，处理作文内容并返回评分结果
    
    Args:
        title (str): 作文标题
        content (str): 作文内容
        
    Returns:
        dict: 包含评分结果的字典
    """
    try:
        # 检查AI服务是否可用
        if not ai_initialized or not ai_available:
            logger.error("AI评分服务不可用，返回备用结果")
            return get_fallback_result(title, content)
            
        # 计算文本长度
        text_length = len(content)
        logger.info(f"开始对作文《{title}》进行AI评分，文本长度：{text_length}字")
        
        # 导入AI功能
        from ai_correction_config import get_unified_prompt
        from tasks import task_unified_scoring
        from llm_func import model
        import json
        
        # 生成统一评分提示词
        prompt = get_unified_prompt(title, content, text_length)
        
        # 格式化为LLM的输入消息格式
        messages = [{"role": "user", "content": prompt}]
        
        # 调用AI模型
        logger.info("调用AI模型进行统一评分")
        response = asyncio.run(model(messages))
        
        # 解析AI返回的结果
        if response and response.choices and len(response.choices) > 0:
            result_text = response.choices[0].message.content
            logger.info(f"AI返回原始结果: {result_text[:200]}...")
            
            # 尝试从返回文本中提取JSON
            try:
                # 查找JSON内容
                json_start = result_text.find('{')
                json_end = result_text.rfind('}') + 1
                
                if json_start >= 0 and json_end > json_start:
                    json_str = result_text[json_start:json_end]
                    result = json.loads(json_str)
                    
                    # 标准化处理AI返回的结果
                    total_score = int(result.get('总得分', '0'))
                    grade = result.get('等级评定', 'C-中等')
                    
                    # 分项得分
                    scores = result.get('分项得分', {})
                    content_score = int(scores.get('内容主旨', '0'))
                    language_score = int(scores.get('语言文采', '0'))
                    structure_score = int(scores.get('文章结构', '0'))
                    writing_score = int(scores.get('文面书写', '0'))
                    
                    # 多维分析
                    analysis = result.get('多维分析', {})
                    
                    # 构建标准格式的返回结果
                    standardized_result = {
                        'total_score': total_score,
                        'grade': grade,
                        'content_score': content_score,
                        'language_score': language_score,
                        'structure_score': structure_score,
                        'writing_score': writing_score,
                        'spelling_errors': result.get('错别字', []),
                        'overall_comment': result.get('总体评价', ''),
                        'detailed_analysis': analysis
                    }
                    
                    logger.info(f"评分完成: 总分={total_score}, 等级={grade}")
                    return standardized_result
                else:
                    logger.error("无法从AI响应中提取JSON数据")
            except json.JSONDecodeError as e:
                logger.error(f"JSON解析错误: {e}")
                logger.error(f"原始文本: {result_text}")
            except Exception as e:
                logger.error(f"处理AI响应时出错: {e}")
        
        # 如果出现任何问题，返回备用结果
        logger.warning("AI评分失败，返回备用结果")
        return get_fallback_result(title, content)
        
    except Exception as e:
        logger.error(f"统一评分函数出错: {str(e)}")
        return get_fallback_result(title, content)

# 获取用户作文批改限制
def get_user_limits(user_type):
    """根据用户类型获取作文批改限制"""
    if user_type == 'free':
        return 2, 10  # 免费用户：每日2篇，每月10篇
    elif user_type == 'regular':
        return 30, 300  # 普通会员：每日30篇，每月300篇
    elif user_type in ['premium', 'vip']:
        return 60, 500  # 高级会员：每日60篇，每月500篇
    else:
        return 2, 10  # 默认使用免费用户限制

# 导入自定义模块
from user_auth import (
    User, register_user, login_user_with_password, login_user_with_code, 
    send_verification_email, save_verification_code, generate_verification_code,
    check_email_verification_limit, reset_password_with_code, reset_password_without_code
)
from user_profile import UserProfile, handle_profile_route
from user_history import EssayHistory, handle_user_history_route, handle_delete_essay_route
# 需要按需导入你的作文批改相关模块
# from essay_correction import ...

# 添加AI评分配置模块导入
from ai_correction_config import (
    get_prompt_template, 
    extract_spelling_errors, 
    format_spelling_errors,
    normalize_scores,
    get_fallback_result,
    SCORING
)

# 导入统一评分任务
from tasks import task_unified_scoring

# 添加路由调试函数
routes_checked = False

@app.before_request
def check_routes():
    global routes_checked
    if not routes_checked:
        logger.info("\n===== 已注册的路由 =====")
        for rule in app.url_map.iter_rules():
            logger.info(f"{rule.endpoint}: {rule}")
        logger.info("=======================\n")
        routes_checked = True

# 通用函数：获取数据库连接
def get_db():
    if 'db' not in g:
        g.db = sqlite3.connect(DATABASE)
        g.db.row_factory = sqlite3.Row
    return g.db

# 通用函数：关闭数据库连接
@app.teardown_appcontext
def close_db(e=None):
    db = g.pop('db', None)
    if db is not None:
        db.close()

# 管理员权限检查装饰器
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('请先登录', 'warning')
            return redirect(url_for('login'))
        
        user_id = session.get('user_id')
        role = session.get('role')
        
        if role != 'admin':
            flash('您没有权限访问此页面', 'danger')
            return redirect(url_for('index'))
            
        return f(*args, **kwargs)
    return decorated_function

# 获取用户最近的作文
def get_recent_essays(user_id, limit=5):
    """获取用户最近批改的作文列表"""
    if not user_id:
        return []
        
    try:
        conn = get_db()
        cursor = conn.cursor()
        
        # 查询用户最近的作文
        cursor.execute("""
            SELECT id, title, submission_time as created_at, total_score
            FROM essays 
            WHERE user_id = ? 
            ORDER BY submission_time DESC
            LIMIT ?
        """, (user_id, limit))
        
        essays = cursor.fetchall()
        
        # 将查询结果转换为字典列表
        result = []
        for essay in essays:
            result.append({
                'id': essay[0],
                'title': essay[1],
                'created_at': essay[2],
                'total_score': essay[3] if essay[3] is not None else 0
            })
        
        return result
    except Exception as e:
        logger.error(f"获取最近作文时出错: {e}")
        return []

# 获取用户统计信息
def get_user_stats(user_id):
    """获取用户作文统计信息"""
    if not user_id:
        return None
        
    try:
        conn = get_db()
        cursor = conn.cursor()
        
        # 查询用户作文总数
        cursor.execute("SELECT COUNT(*) FROM essays WHERE user_id = ?", (user_id,))
        total_essays = cursor.fetchone()[0]
        
        # 查询用户平均分数（排除NULL值）
        cursor.execute("""
            SELECT ROUND(AVG(total_score), 2) as avg_score
            FROM essays 
            WHERE user_id = ? AND total_score IS NOT NULL
        """, (user_id,))
        avg_score_result = cursor.fetchone()[0]
        avg_score = avg_score_result if avg_score_result is not None else None
        
        # 获取最高分和总篇数
        cursor.execute("""
            SELECT MAX(total_score) as max_score, COUNT(*) as essay_count
            FROM essays
            WHERE user_id = ? AND total_score IS NOT NULL
        """, (user_id,))
        stats_row = cursor.fetchone()
        max_score = stats_row[0]
        essay_count = stats_row[1]
        
        return {
            'total_essays': total_essays,
            'avg_score': avg_score,
            'max_score': max_score,
            'essay_count': essay_count
        }
    except Exception as e:
        logger.error(f"获取用户统计信息时出错: {e}")
        return None

# 主页路由
@app.route('/')
def index():
    # 检查用户是否已登录
    user = None
    logged_in = False
    recent_essays = []
    stats = None
    
    if 'user_id' in session:
        logged_in = True
        user_id = session['user_id']
        user = User.get(user_id)
        
        # 如果需要，这里可以获取最近的作文和统计信息
        recent_essays = get_recent_essays(user_id)
        # stats = get_user_stats(user_id)
    
    # 渲染首页模板
    return render_template('index.html', 
                        user=user, 
                        logged_in=logged_in,
                        recent_essays=recent_essays,
                        stats=stats)

# 用户作文历史路由
@app.route('/user_history')
def user_history():
    """用户作文历史页面"""
    # 检查用户是否已登录
    if 'user_id' not in session:
        flash('请先登录后再查看作文历史', 'warning')
        return redirect(url_for('login'))
    
    return handle_user_history_route()

# 登录页面路由
@app.route('/login', methods=['GET', 'POST'])
def login():
    # 如果用户已登录，重定向到主页
    if 'user_id' in session:
        return redirect(url_for('index'))
        
    if request.method == 'POST':
        # 使用密码登录
        username = request.form.get('username')
        password = request.form.get('password')
        remember = request.form.get('remember', '0') == '1'
        
        success, result = login_user_with_password(username, password)
        
        if success:
            # 登录成功，设置session
            user = result
            session['user_id'] = user['user_id']
            session['role'] = user.get('role', 'user')
            if remember:
                session.permanent = True
            
            # 更新最后登录时间
            profile_manager = UserProfile()
            profile_manager.update_last_login(user['user_id'])
            
            flash('登录成功', 'success')
            return redirect(url_for('index'))
        else:
            flash(result, 'danger')
    
    return render_template('login.html')

# 退出登录路由
@app.route('/logout')
def logout():
    session.clear()
    flash('您已成功退出登录', 'info')
    return redirect(url_for('index'))

# 忘记密码路由
@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    """忘记密码页面，支持通过电子邮件重置密码"""
    if request.method == 'POST':
        email = request.form.get('email')
        if not email:
            flash('请输入邮箱地址', 'danger')
            return render_template('forgot_password.html')
            
        try:
            # 检查邮箱是否存在
            conn = get_db()
            cursor = conn.cursor()
            cursor.execute("SELECT user_id FROM users WHERE email = ?", (email,))
            user = cursor.fetchone()
            
            if not user:
                flash('该邮箱未注册', 'danger')
                return render_template('forgot_password.html')
                
            # 生成验证码
            verification_code = generate_verification_code()
            
            # 发送验证邮件
            send_verification_email(email, verification_code, 'reset_password')
            
            # 保存验证码
            save_verification_code(email, verification_code, 'reset_password')
            
            flash('重置密码的验证码已发送到您的邮箱，请查收', 'success')
            return redirect(url_for('reset_password'))
        except Exception as e:
            logger.error(f"发送重置密码验证码时出错: {str(e)}")
            flash('发送验证码失败，请稍后重试', 'danger')
    
    return render_template('forgot_password.html')

# 重置密码路由
@app.route('/reset_password', methods=['GET', 'POST'])
def reset_password():
    """重置密码页面"""
    if request.method == 'POST':
        email = request.form.get('email')
        code = request.form.get('code')
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')
        
        if not all([email, code, new_password, confirm_password]):
            flash('请填写所有必填字段', 'danger')
            return render_template('reset_password.html')
            
        if new_password != confirm_password:
            flash('两次输入的密码不一致', 'danger')
            return render_template('reset_password.html')
            
        success, message = reset_password_with_code(email, code, new_password)
        
        if success:
            flash('密码重置成功，请使用新密码登录', 'success')
            return redirect(url_for('login'))
        else:
            flash(message, 'danger')
            return render_template('reset_password.html')
    
    return render_template('reset_password.html')

# 注册路由
@app.route('/register', methods=['GET', 'POST'])
def register():
    """用户注册页面"""
    # 如果用户已登录，重定向到主页
    if 'user_id' in session:
        return redirect(url_for('index'))
        
    if request.method == 'POST':
        # 获取表单数据
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        verification_code = request.form.get('verification_code')
        
        # 基本验证
        if not all([username, email, password, confirm_password]):
            flash('请填写所有必填字段', 'danger')
            return render_template('register.html')
            
        if password != confirm_password:
            flash('两次输入的密码不一致', 'danger')
            return render_template('register.html')
            
        try:
            # 检查用户名是否已存在
            conn = get_db()
            cursor = conn.cursor()
            cursor.execute("SELECT 1 FROM users WHERE username = ?", (username,))
            if cursor.fetchone():
                flash('用户名已被使用', 'danger')
                return render_template('register.html')
                
            # 检查邮箱是否已存在
            cursor.execute("SELECT 1 FROM users WHERE email = ?", (email,))
            if cursor.fetchone():
                flash('邮箱已被注册', 'danger')
                return render_template('register.html')
                
            # 验证验证码
            if verification_code:
                is_verified = verify_code(email, verification_code, 'register')
                if not is_verified:
                    flash('验证码无效或已过期', 'danger')
                    return render_template('register.html')
            
            # 创建用户
            password_hash = generate_password_hash(password)
            cursor.execute("""
                INSERT INTO users (username, email, password_hash, created_at, user_type, role, is_active)
                VALUES (?, ?, ?, datetime('now'), 'free', 'user', 1)
            """, (username, email, password_hash))
            conn.commit()
            
            # 获取用户ID
            cursor.execute("SELECT user_id FROM users WHERE username = ?", (username,))
            user_id = cursor.fetchone()[0]
            
            # 设置session
            session['user_id'] = user_id
            session['role'] = 'user'
            
            flash('注册成功，欢迎加入！', 'success')
            return redirect(url_for('index'))
            
        except Exception as e:
            logger.error(f"注册用户时出错: {str(e)}")
            flash('注册失败，请稍后重试', 'danger')
            
    return render_template('register.html')

# 个人资料路由
@app.route('/profile', methods=['GET', 'POST'])
def profile():
    """用户个人资料页面"""
    # 检查用户是否已登录
    if 'user_id' not in session:
        flash('请先登录后再访问个人资料', 'warning')
        return redirect(url_for('login'))
        
    user_id = session['user_id']
    
    try:
        # 获取用户信息
        conn = get_db()
        cursor = conn.cursor()
        
        # 获取用户基本信息
        cursor.execute("SELECT * FROM users WHERE user_id = ?", (user_id,))
        user_data = cursor.fetchone()
        
        if not user_data:
            flash('用户不存在', 'danger')
            return redirect(url_for('index'))
            
        # 将结果转换为字典
        user = dict(zip([column[0] for column in cursor.description], user_data))
        
        # 获取用户资料详情
        cursor.execute("SELECT * FROM user_profiles WHERE user_id = ?", (user_id,))
        profile_data = cursor.fetchone()
        
        # 如果资料不存在，创建空资料
        if not profile_data:
            profile = {
                'user_id': user_id,
                'full_name': '',
                'school': '',
                'grade': '',
                'created_at': user['created_at'],
                'last_login': None
            }
        else:
            # 将结果转换为字典
            profile = dict(zip([column[0] for column in cursor.description], profile_data))
        
        if request.method == 'POST':
            form_type = request.form.get('form_type')
            
            if form_type == 'profile':
                # 更新个人资料
                full_name = request.form.get('full_name', '')
                school = request.form.get('school', '')
                grade = request.form.get('grade', '')
                
                # 更新或插入资料
                if profile_data:
                    cursor.execute("""
                        UPDATE user_profiles 
                        SET full_name = ?, school = ?, grade = ? 
                        WHERE user_id = ?
                    """, (full_name, school, grade, user_id))
                else:
                    cursor.execute("""
                        INSERT INTO user_profiles (user_id, full_name, school, grade)
                        VALUES (?, ?, ?, ?)
                    """, (user_id, full_name, school, grade))
                
                conn.commit()
                flash('个人资料更新成功', 'success')
                
                # 更新资料后重新获取
                profile['full_name'] = full_name
                profile['school'] = school
                profile['grade'] = grade
                
            elif form_type == 'password':
                # 修改密码
                current_password = request.form.get('current_password')
                new_password = request.form.get('new_password')
                confirm_password = request.form.get('confirm_password')
                
                if not all([current_password, new_password, confirm_password]):
                    flash('请填写所有密码字段', 'danger')
                    return render_template('profile.html', user=user, profile=profile)
                    
                if new_password != confirm_password:
                    flash('两次输入的新密码不一致', 'danger')
                    return render_template('profile.html', user=user, profile=profile)
                
                # 验证当前密码
                if check_password_hash(user['password_hash'], current_password):
                    # 更新密码
                    new_password_hash = generate_password_hash(new_password)
                    cursor.execute("UPDATE users SET password_hash = ? WHERE user_id = ?", 
                                  (new_password_hash, user_id))
                    conn.commit()
                    flash('密码修改成功', 'success')
                else:
                    flash('当前密码不正确', 'danger')
        
        return render_template('profile.html', user=user, profile=profile)
        
    except Exception as e:
        logger.error(f"加载个人资料页面出错: {str(e)}")
        flash(f'加载个人资料页面出错: {str(e)}', 'danger')
        return redirect(url_for('index'))

# 下面的所有代码是路由定义，放在app创建之后，确保路由注册到正确的app实例
# 主页路由
@app.route('/correction', methods=['GET', 'POST'])
def correction():
    """作文批改页面：接收用户提交的作文并进行批改"""
    if request.method == 'POST':
        try:
            # 检查用户登录状态
            user_id = session.get('user_id')
            if not user_id:
                flash('请先登录后再进行作文批改', 'warning')
                return redirect(url_for('login'))
    
            # 检查用户剩余次数
            connection = get_db()
            cursor = connection.cursor()
            
            # 获取用户信息和会员状态
            cursor.execute("SELECT user_type, membership_expiry FROM users WHERE user_id = ?", (user_id,))
            user_info = cursor.fetchone()
            
            # 如果用户不存在，给出错误提示
            if not user_info:
                flash('用户信息无效，请重新登录', 'danger')
                return redirect(url_for('login'))
    
            user_type = user_info[0]
            membership_expiry = user_info[1]
            
            # 检查会员是否有效
            is_premium = False
            if membership_expiry:
                try:
                    expiry_date = datetime.strptime(membership_expiry, '%Y-%m-%d').date()
                    is_premium = expiry_date >= date.today() and user_type in ['regular', 'premium']
                    logger.info(f"用户ID {user_id} 会员状态: {user_type}, 到期日: {expiry_date}, 当前有效: {is_premium}")
                except ValueError as e:
                    logger.error(f"会员到期日期解析错误: {e}")
                    flash('系统错误: 会员信息异常', 'danger')
                    return redirect(url_for('index'))
                    
            # 获取当天已批改的作文数量
            today = date.today().strftime('%Y-%m-%d')
            cursor.execute("""
                SELECT COUNT(*) FROM essays 
                WHERE user_id = ? AND DATE(submission_time) = ?
            """, (user_id, today))
            daily_count = cursor.fetchone()[0]
            
            # 获取当月已批改的作文数量
            current_month = date.today().strftime('%Y-%m')
            cursor.execute("""
                SELECT COUNT(*) FROM essays 
                WHERE user_id = ? AND strftime('%Y-%m', submission_time) = ?
            """, (user_id, current_month))
            monthly_count = cursor.fetchone()[0]
            
            # 检查剩余量
            daily_limit, monthly_limit = get_user_limits(user_type)
            logger.info(f"用户ID {user_id} 限制: 每日{daily_limit}, 每月{monthly_limit}, 已用: 今日{daily_count}, 本月{monthly_count}")
            
            # 检查是否超出限制
            if daily_count >= daily_limit:
                flash(f'您今日批改已达上限({daily_limit}篇)，请明天再试或升级会员', 'warning')
                return redirect(url_for('membership'))
            
            if monthly_count >= monthly_limit:
                flash(f'您本月批改已达上限({monthly_limit}篇)，请下月再试或升级会员', 'warning')
                return redirect(url_for('membership'))
            
            # 获取表单数据
            subject = request.form.get('subject', '无标题').strip()
            article = request.form.get('article', '').strip()
            file = request.files.get('file')
            use_file_content = request.form.get('useFileContent') == 'on'
            
            # 如果选择了使用文件内容且有上传文件
            if file and file.filename:
                file_content = None
                file_extension = os.path.splitext(file.filename)[1].lower()

                try:
                    # 根据文件扩展名处理不同格式
                    if file_extension == '.txt':
                        file_content = file.read().decode('utf-8')
                    elif file_extension == '.docx':
                        import docx
                        doc = docx.Document(file)
                        file_content = '\n'.join([para.text for para in doc.paragraphs])
                    else:
                        flash('不支持的文件格式，请上传.txt或.docx文件', 'danger')
                        return redirect(url_for('correction'))
                    
                    # 如果成功获取了文件内容
                    if file_content:
                        article = file_content
                        # 如果没有提供标题，尝试使用文件名作为标题
                        if not subject:
                            subject = os.path.splitext(file.filename)[0]
                except Exception as e:
                    logger.error(f"文件处理出错: {str(e)}")
                    flash(f'文件读取失败: {str(e)}', 'danger')
                    return redirect(url_for('correction'))
            
            # 验证内容是否为空
            if not article.strip():
                flash('作文内容不能为空，请输入作文内容', 'danger')
                return redirect(url_for('correction'))
            
            # 如果没有标题，尝试从内容的第一行提取
            if not subject:
                first_line = article.strip().split('\n')[0].strip()
                if len(first_line) <= 30:  # 如果第一行不太长，可能是标题
                    subject = first_line
                else:
                    subject = '无标题'
            
            # 获取字数
            word_count = len(article)
            
            # 记录到数据库
            submission_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            try:
                cursor.execute("""
                    INSERT INTO essays 
                    (user_id, title, content, submission_time, word_count) 
                    VALUES (?, ?, ?, ?, ?)
                """, (user_id, subject, article, submission_time, word_count))
                
                essay_id = cursor.lastrowid
                connection.commit()
                
                # 调用AI评分服务
                try:
                    ai_result = unified_ai_scoring(subject, article)
                    
                    # 确保等级正确，特别是对空作文
                    grade = ai_result.get('grade', 'E-未完成')
                    if len(article.strip()) == 0 or int(ai_result.get('total_score', 0)) == 0:
                        grade = "E-未完成"
                    
                    # 获取错别字数组
                    spelling_errors = json.dumps(ai_result.get('spelling_errors', []), ensure_ascii=False)
                    
                    # 更新评分结果到数据库
                    cursor.execute("""
                        UPDATE essays SET 
                            total_score = ?, grade = ?, content_score = ?, 
                            language_score = ?, structure_score = ?, writing_score = ?,
                            spelling_errors = ?, overall_assessment = ?, 
                            content_analysis = ?, language_analysis = ?,
                            structure_analysis = ?, writing_analysis = ?
                        WHERE id = ?
                    """, (
                        ai_result.get('total_score', 0),
                        grade,
                        ai_result.get('content_score', 0),
                        ai_result.get('language_score', 0),
                        ai_result.get('structure_score', 0),
                        ai_result.get('writing_score', 0),
                        spelling_errors,
                        ai_result.get('overall_comment', ''),
                        ai_result.get('detailed_analysis', {}).get('内容分析', ''),
                        ai_result.get('detailed_analysis', {}).get('表达分析', ''),
                        ai_result.get('detailed_analysis', {}).get('结构分析', ''),
                        ai_result.get('detailed_analysis', {}).get('书写分析', ''),
                        essay_id
                    ))
                    connection.commit()
                    
                    # 重定向到结果页面
                    return redirect(url_for('results', essay_id=essay_id))
                    
                except Exception as ai_error:
                    logger.error(f"AI评分出错: {str(ai_error)}")
                    flash(f'AI评分出错: {str(ai_error)}', 'danger')
                    return redirect(url_for('correction'))
                    
            except Exception as db_error:
                logger.error(f"数据库操作出错: {str(db_error)}")
                flash(f'保存作文出错: {str(db_error)}', 'danger')
                return redirect(url_for('correction'))
                
        except Exception as e:
            logger.error(f"发生错误: {str(e)}")
            flash('系统错误，请稍后再试', 'danger')
            return redirect(url_for('index'))
    
    # GET请求直接返回批改页面
    return render_template('correction.html')

@app.route('/results')
def results():
    """展示作文的批改结果"""
    try:
        essay_id = request.args.get('essay_id')
        
        # 获取数据库连接
        conn = get_db()
        cursor = conn.cursor()
        
        # 如果提供了文章ID，从数据库获取
        if essay_id:
            cursor.execute("""
                SELECT * FROM essays WHERE id = ?
            """, (essay_id,))
            essay = cursor.fetchone()
            
            if not essay:
                flash('未找到指定的作文', 'warning')
                return redirect(url_for('correction'))
            
            # 将数据库结果转为字典
            essay_dict = {key: essay[key] for key in essay.keys()}
            
            # 处理错别字信息
            spelling_errors = []
            if essay['spelling_errors']:
                try:
                    spelling_errors = json.loads(essay['spelling_errors'])
                except Exception as e:
                    logger.error(f"解析错别字数据出错: {str(e)}")
            
            # 准备模板渲染数据
            template_data = {
                'title': essay['title'],
                'content': essay['content'],
                'total_score': essay['total_score'],
                'grade': essay['grade'],
                'content_score': essay['content_score'],
                'language_score': essay['language_score'],
                'structure_score': essay['structure_score'],
                'writing_score': essay['writing_score'],
                'spelling_errors': spelling_errors,
                'word_count': essay['word_count'],
                'overall_assessment': essay['overall_assessment'],
                'content_analysis': essay['content_analysis'],
                'language_analysis': essay['language_analysis'],
                'structure_analysis': essay['structure_analysis'],
                'writing_analysis': essay['writing_analysis']
            }
        else:
            flash('未提供作文ID', 'warning')
            return redirect(url_for('correction'))
        
        return render_template('results.html', **template_data)
    except Exception as e:
        logger.error(f"显示作文结果时出错: {str(e)}")
        flash(f'显示作文结果时出错: {str(e)}', 'danger')
        return redirect(url_for('correction'))

# 批量上传作文路由
@app.route('/batch_upload', methods=['GET', 'POST'])
def batch_upload():
    # 检查用户是否已登录
    if 'user_id' not in session:
        flash('请先登录后再进行批量上传', 'warning')
        return redirect(url_for('login'))
    
    processed_files_info = [] # 用于存储每个文件的处理结果
    
    if request.method == 'POST':
        if 'files[]' not in request.files:
            flash('未检测到上传的文件', 'danger')
            return redirect(request.url)
        
        files = request.files.getlist('files[]')
        
        if not files or len(files) == 0 or files[0].filename == '':
            flash('请选择要上传的文件', 'danger')
            return redirect(request.url)
        
        user_id = session.get('user_id')
        conn = None # 初始化数据库连接
        
        try:
            conn = sqlite3.connect('instance/essay_correction.db')
            cursor = conn.cursor()
            
            for file in files:
                filename = file.filename
                content = ""
                status = '失败' # 默认状态
                error_message = ''
                essay_id = None
                
                try:
                    filename_base = os.path.splitext(filename)[0]
                    file_ext = os.path.splitext(filename)[1].lower()
                    
                    # 读取文件内容
                    if file_ext == '.txt':
                        content = file.read().decode('utf-8')
                    elif file_ext == '.docx':
                        import docx
                        doc = docx.Document(file)
                        content = '\n'.join([para.text for para in doc.paragraphs])
                    elif file_ext == '.doc':
                        try:
                            import pypandoc
                            # Ensure upload folder exists
                            upload_dir = app.config.get('UPLOAD_FOLDER', 'uploads')
                            if not os.path.exists(upload_dir):
                                os.makedirs(upload_dir)
                            # Save the uploaded file temporarily to read with pandoc
                            temp_filename = os.path.join(upload_dir, secure_filename(f"{uuid.uuid4()}_{filename}")) # Add UUID
                            file.seek(0) # Reset file pointer before saving
                            file.save(temp_filename)
                            # Convert .doc to plain text using pandoc
                            content = pypandoc.convert_file(temp_filename, 'plain', format='doc')
                            os.remove(temp_filename) # Clean up the temporary file
                            logger.info(f"Successfully read .doc file '{filename}' using pandoc")
                        except ImportError:
                            logger.error("pypandoc library not found. Please install it: pip install pypandoc")
                            error_message = '无法处理.doc文件，缺少pypandoc库'
                            raise ValueError(error_message)
                        except Exception as pandoc_error:
                            logger.error(f"Error reading .doc file '{filename}' with pandoc: {pandoc_error}")
                            error_message = f'处理.doc文件 ({filename}) 时出错: {pandoc_error}'
                            if os.path.exists(temp_filename): # Clean up temp file on error too
                                os.remove(temp_filename)
                            raise ValueError(error_message) # Raise error to be caught by outer loop
                    else:
                        error_message = f"不支持的文件格式: {file_ext}"
                        raise ValueError(error_message)
                        
                    # 检查内容是否为空
                    if not content.strip():
                        error_message = "文件内容为空"
                        raise ValueError(error_message)
                    
                    # 获取标题
                    title = filename_base # 批量上传优先使用文件名
                    logger.info(f"处理文件: '{filename}', 获取标题: '{title}'")

                    # 1. 先插入基本信息到数据库
                    current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    word_count = len(content)
                    cursor.execute(
                        "INSERT INTO essays (user_id, title, content, submission_time, word_count) VALUES (?, ?, ?, ?, ?)",
                        (user_id, title, content, current_time, word_count)
                    )
                    essay_id = cursor.lastrowid
                    conn.commit() # 提交事务以获取 ID
                    logger.info(f"文件 '{filename}' 基本信息已存入数据库, ID={essay_id}")
                        
                    # 2. 调用AI进行评分
                    logger.info(f"开始对文件 '{filename}' (ID: {essay_id}) 进行AI评分")
                    ai_result = unified_ai_scoring(title, content)
                    logger.info(f"文件 '{filename}' AI评分完成")

                    # 调试输出，确保所有评分结果正确获取
                    logger.info(f"AI评分详细结果: 总分={ai_result.get('总分', 0)}, 等级={ai_result.get('等级', 'E')}")
                    logger.info(f"多维分析内容: 内容分析={ai_result.get('内容分析', '')[:50]}..., 表达分析={ai_result.get('表达分析', '')[:50]}...")
                    logger.info(f"总体评价: {ai_result.get('总体评价', '')[:50]}...")

                    # 3. 更新数据库中的评分结果
                    spelling_errors_json = json.dumps(ai_result.get('错别字', {"解析": []}), ensure_ascii=False)
                    grade = ai_result.get('grade_assessment', ai_result.get('grade', 'E-未完成')) # 获取等级，优先使用grade_assessment
                    
                    cursor.execute("""
                        UPDATE essays SET 
                        total_score = ?, grade = ?, content_score = ?, 
                        language_score = ?, structure_score = ?, writing_score = ?,
                        spelling_errors = ?, overall_assessment = ?, 
                        content_analysis = ?, language_analysis = ?,
                        structure_analysis = ?, writing_analysis = ?
                        WHERE id = ?
                    """, (
                        ai_result.get('total_score', 0),
                        grade,
                        ai_result.get('content_score', 0),
                        ai_result.get('language_score', 0),
                        ai_result.get('structure_score', 0),
                        ai_result.get('writing_score', 0),
                        spelling_errors_json,
                        ai_result.get('overall_comment', ''),
                        ai_result.get('detailed_analysis', {}).get('内容分析', ''),
                        ai_result.get('detailed_analysis', {}).get('表达分析', ''),
                        ai_result.get('detailed_analysis', {}).get('结构分析', ''),
                        ai_result.get('detailed_analysis', {}).get('书写分析', ''),
                        essay_id
                    ))
                    conn.commit() # 提交评分更新
                    logger.info(f"文件 '{filename}' (ID: {essay_id}) 评分结果已更新至数据库")
                    status = '成功'
                        
                except Exception as e:
                    logger.error(f"处理文件 '{filename}' 时出错: {str(e)}")
                    error_message = str(e)
                    if conn: # 如果出错，回滚当前文件的事务
                        conn.rollback()
                
                finally:
                    # 记录处理结果
                    processed_files_info.append({
                        'filename': filename,
                        'status': status,
                        'error': error_message,
                        'essay_id': essay_id # 可以用来生成结果链接
                    })
                    file.seek(0) # 重置文件指针以便读取大小（如果需要）

            flash(f'批量上传处理完成，共处理 {len(files)} 个文件', 'info')
                
        except Exception as e:
            logger.error(f"批量上传过程中发生严重错误: {str(e)}")
            flash('批量上传过程中发生错误，部分文件可能未处理', 'danger')
            if conn:
                conn.rollback() # 回滚所有未提交的更改
                
        finally:
            if conn:
                conn.close() # 确保数据库连接被关闭
        
        # 渲染结果，显示每个文件的处理状态
        # 保存成功处理的文件ID到session，用于在用户历史页面高亮显示
        successful_essay_ids = [file_info.get('essay_id') for file_info in processed_files_info if file_info.get('status') == '成功' and file_info.get('essay_id')]
        if successful_essay_ids:
            session['last_batch_essay_id'] = successful_essay_ids
            
        return render_template('batch_upload.html', processed_files=processed_files_info)
    
    # GET 请求
    return render_template('batch_upload.html', processed_files=processed_files_info)

@app.route('/membership')
def membership():
    """会员中心页面，显示用户的会员信息和订单历史"""
    # 检查用户是否已登录
    if 'user_id' not in session:
        flash('请先登录后再访问会员中心', 'warning')
        return redirect(url_for('login'))
    
    user_id = session.get('user_id')
    
    # 获取用户信息
    conn = sqlite3.connect('instance/essay_correction.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # 查询用户基本信息
    cursor.execute("SELECT user_id, username, email, user_type, membership_expiry FROM users WHERE user_id = ?", (user_id,))
    user = dict(cursor.fetchone())
    
    # 检查用户是否存在
    if not user:
        conn.close()
        flash('用户信息不存在', 'danger')
        return redirect(url_for('index'))
    
    # 获取会员套餐信息
    membership_plans = []
    cursor.execute("SELECT * FROM membership_plans WHERE active = 1 ORDER BY price")
    for row in cursor.fetchall():
        membership_plans.append(dict(row))
    
    # 获取最近的会员订单
    recent_orders = []
    try:
        # 获取表结构以确定正确的列名
        cursor.execute("PRAGMA table_info(membership_orders)")
        columns = [column[1] for column in cursor.fetchall()]
        
        # 构建查询
        select_columns = []
        if 'id' in columns:
            select_columns.append('id')
        else:
            # 使用主键列名
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='membership_orders'")
            if cursor.fetchone():
                cursor.execute("PRAGMA table_info(membership_orders)")
                pk_columns = [column[1] for column in cursor.fetchall() if column[5] == 1]  # 第6列是pk标志
                if pk_columns:
                    select_columns.append(f"{pk_columns[0]} as id")
        
        # 添加其他列
        for col in ['plan_id', 'plan_name', 'amount', 'created_at', 'payment_status']:
            if col in columns:
                select_columns.append(col)
        
        if select_columns:
            query = f"""
                SELECT {', '.join(select_columns)}
                FROM membership_orders 
                WHERE user_id = ? 
                ORDER BY created_at DESC 
                LIMIT 5
            """
            cursor.execute(query, (user_id,))
            for row in cursor.fetchall():
                recent_orders.append(dict(row))
    except Exception as e:
        logger.error(f"查询会员订单时出错: {str(e)}")
        # 出错时使用空列表
    
    # 获取用户统计信息 - 已批改的作文数量等
    cursor.execute("SELECT COUNT(*) as essay_count FROM essays WHERE user_id = ?", (user_id,))
    essay_count = cursor.fetchone()['essay_count']
    
    # 添加模板所需的其他用户信息
    user['essays_total_used'] = essay_count
    
    # 根据用户类型设置不同的默认剩余次数
    if user['user_type'] == 'free':
        user['essays_remaining'] = 10  # 免费用户10篇
        user['essays_daily_limit'] = 2  # 免费用户每日限制
    elif user['user_type'] == 'regular':
        user['essays_remaining'] = 300  # 普通会员300篇
        user['essays_daily_limit'] = 30  # 普通会员每日限制
    elif user['user_type'] in ['premium', 'vip']:
        user['essays_remaining'] = 500  # 高级会员500篇
        user['essays_daily_limit'] = 60  # 高级会员每日限制
    else:
        user['essays_remaining'] = 10  # 默认为免费用户权限
        user['essays_daily_limit'] = 2  # 默认每日限制
    
    user['essays_daily_used'] = 0  # 今日已使用次数
    user['vip_status'] = 0  # 非无限状态
    
    # 配置信息供模板使用，与用户类型权限保持一致
    config = {
        'free_essays': 10,              # 免费用户总批改次数
        'free_daily_limit': 2,          # 免费用户每日限制
        'regular_monthly_essays': 300,  # 普通会员批改次数
        'regular_daily_essays': 30,     # 普通会员每日限制
        'premium_monthly_essays': 500,  # 高级会员批改次数
        'premium_daily_essays': 60      # 高级会员每日限制
    }
    
    conn.close()
    
    # 处理会员过期时间
    membership_status = '普通用户'
    days_left = 0
    
    if user['user_type'] in ['premium', 'vip']:
        membership_status = '会员用户'
        # 计算剩余天数
        if user['membership_expiry']:
            try:
                # 尝试完整的日期时间格式
                expiry_date = datetime.strptime(user['membership_expiry'], '%Y-%m-%d %H:%M:%S')
            except ValueError:
                try:
                    # 尝试仅日期格式
                    expiry_date = datetime.strptime(user['membership_expiry'], '%Y-%m-%d')
                except ValueError:
                    # 记录无法解析的日期格式
                    logger.error(f"无法解析的会员过期日期格式: {user['membership_expiry']}")
                    expiry_date = datetime.now()
            
            days_left = (expiry_date - datetime.now()).days
            if days_left < 0:
                days_left = 0
                membership_status = '会员已过期'
    
    # 渲染会员中心页面
    return render_template('membership.html', 
                        user=user,
                        membership_status=membership_status,
                        days_left=days_left,
                        essay_count=essay_count,
                        recent_orders=recent_orders,
                        membership_plans=membership_plans,
                        config=config)

@app.route('/purchase_membership')
@app.route('/purchase_membership/<string:plan>')
def purchase_membership(plan=None):
    """购买会员页面"""
    # 检查用户是否已登录
    if 'user_id' not in session:
        flash('请先登录后再购买会员', 'warning')
        return redirect(url_for('login'))
    
    # 简单的重定向回会员中心，实际项目中此处应该是支付流程
    flash('会员购买功能正在开发中，敬请期待', 'info')
    return redirect(url_for('membership'))
    
@app.route('/order_history')
def order_history():
    """会员订单历史页面，显示用户的订单记录"""
    # 检查用户是否已登录
    if 'user_id' not in session:
        flash('请先登录后再查看订单历史', 'warning')
        return redirect(url_for('login'))
    
    user_id = session.get('user_id')
    
    # 获取分页参数
    page = request.args.get('page', 1, type=int)
    per_page = 10  # 每页显示10条记录
    
    # 查询用户的订单
    conn = sqlite3.connect('instance/essay_correction.db')
    conn.row_factory = sqlite3.Row  # 使结果可以通过列名访问
    cursor = conn.cursor()
    
    # 获取总记录数
    total_orders = 0
    orders = []
    
    try:
        # 获取总记录数
        cursor.execute("SELECT COUNT(*) FROM membership_orders WHERE user_id = ?", (user_id,))
        total_orders = cursor.fetchone()[0]
        
        # 计算分页信息
        total_pages = (total_orders + per_page - 1) // per_page  # 向上取整
        offset = (page - 1) * per_page
        
        # 查询当前页的订单
        cursor.execute("""
            SELECT * FROM membership_orders 
            WHERE user_id = ? 
            ORDER BY created_at DESC
            LIMIT ? OFFSET ?
        """, (user_id, per_page, offset))
        
        for row in cursor.fetchall():
            orders.append(dict(row))
    except Exception as e:
        logger.error(f"查询订单历史时出错: {str(e)}")
        flash('获取订单历史时出错', 'warning')
        total_orders = 0
        total_pages = 0
    
    conn.close()
    
    # 构造分页信息
    pagination = {
        'current_page': page,
        'pages': total_pages,
        'total': total_orders
    }
    
    return render_template('order_history.html', orders=orders, pagination=pagination)

# ====================== 管理后台路由 ======================

# 管理控制台首页
@app.route('/admin/dashboard')
@admin_required
def admin_dashboard():
    """管理后台控制面板"""
    try:
        # 获取数据库连接
        conn = get_db()
        cursor = conn.cursor()
        
        # 获取统计数据
        cursor.execute("SELECT COUNT(*) FROM users")
        total_users = cursor.fetchone()[0]
        
        cursor.execute("SELECT COUNT(*) FROM essays")
        total_essays = cursor.fetchone()[0]
        
        cursor.execute("SELECT COUNT(*) FROM users WHERE user_type IN ('regular', 'premium')")
        premium_users = cursor.fetchone()[0]
        
        # 获取总收入
        cursor.execute("SELECT SUM(amount) FROM membership_orders WHERE payment_status = 'completed'")
        total_income = cursor.fetchone()[0] or 0
        
        # 获取月度收入统计
        cursor.execute("""
            SELECT 
                strftime('%Y-%m', created_at) as month,
                SUM(amount) as monthly_income,
                COUNT(*) as order_count
            FROM membership_orders 
            WHERE payment_status = 'completed'
            GROUP BY month
            ORDER BY month DESC
            LIMIT 6
        """)
        monthly_stats = []
        for row in cursor.fetchall():
            monthly_stats.append({
                'month': row[0], 
                'income': row[1] or 0, 
                'count': row[2] or 0
            })
        
        # 获取最近注册的用户
        cursor.execute("""
            SELECT user_id, username, email, created_at, user_type 
            FROM users 
            ORDER BY created_at DESC 
            LIMIT 5
        """)
        recent_users = [dict(zip([column[0] for column in cursor.description], row)) for row in cursor.fetchall()]
        
        # 获取最近批改的作文
        cursor.execute("""
            SELECT e.id, e.title, e.submission_time, e.total_score, u.username
            FROM essays e
            JOIN users u ON e.user_id = u.user_id
            ORDER BY e.submission_time DESC
            LIMIT 5
        """)
        recent_essays = [dict(zip([column[0] for column in cursor.description], row)) for row in cursor.fetchall()]
        
        # 获取系统信息
        system_info = {
            'ai_available': ai_available,
            'uptime': '正常运行',
            'error_logs': 0,
            'last_backup': None
        }
        
        # 构建统计数据
        stats = {
            'total_users': total_users,
            'total_essays': total_essays,
            'premium_users': premium_users,
            'total_income': round(total_income, 2),
            'monthly_stats': monthly_stats
        }
        
        return render_template('admin/dashboard.html', 
                            stats=stats, 
                            recent_users=recent_users,
                            recent_essays=recent_essays,
                            system_info=system_info,
                            active_page='dashboard')
    except Exception as e:
        logger.error(f"加载管理控制台出错: {str(e)}")
        flash(f'加载管理控制台出错: {str(e)}', 'danger')
        return redirect(url_for('index'))

# 用户管理页面
@app.route('/admin/users')
@admin_required
def user_management():
    """用户管理页面"""
    try:
        # 获取查询参数
        search = request.args.get('search', '')
        user_type = request.args.get('user_type', '')
        role = request.args.get('role', '')
        status = request.args.get('status', '')
        page = request.args.get('page', 1, type=int)
        per_page = 10 # 每页显示10条记录
        
        # 构建SQL查询
        query = "SELECT * FROM users WHERE 1=1"
        params = []
        
        if search:
            query += " AND (username LIKE ? OR email LIKE ?)"
            params.extend([f'%{search}%', f'%{search}%'])
        
        if user_type:
            query += " AND user_type = ?"
            params.append(user_type)
        
        if role:
            query += " AND role = ?"
            params.append(role)
        
        if status:
            query += " AND is_active = ?"
            params.append(int(status))
        
        # 获取数据库连接
        conn = get_db()
        cursor = conn.cursor()
        
        # 获取总记录数
        count_query = f"SELECT COUNT(*) FROM ({query})"
        cursor.execute(count_query, params)
        total = cursor.fetchone()[0]
        
        # 计算分页信息
        total_pages = (total + per_page - 1) // per_page  # 向上取整
        offset = (page - 1) * per_page
        
        # 获取当前页的数据
        query += " ORDER BY created_at DESC LIMIT ? OFFSET ?"
        params.extend([per_page, offset])
        
        cursor.execute(query, params)
        rows = cursor.fetchall()
        
        # 将结果转换为字典列表
        users = []
        for row in rows:
            user = dict(zip([column[0] for column in cursor.description], row))
            users.append(user)
        
        # 构造分页信息
        pagination = {
            'current_page': page,
            'pages': total_pages,
            'total': total
        }
        
        return render_template('admin/user_management.html', 
                            users=users, 
                            pagination=pagination,
                            active_page='users')
    except Exception as e:
        logger.error(f"加载用户管理页面出错: {str(e)}")
        flash(f'加载用户管理页面出错: {str(e)}', 'danger')
        return redirect(url_for('admin_dashboard'))

# 用户详情页面
@app.route('/admin/users/<int:user_id>')
@admin_required
def user_detail(user_id):
    """用户详情页面"""
    try:
        # 获取数据库连接
        conn = get_db()
        cursor = conn.cursor()
        
        # 获取用户信息
        cursor.execute("SELECT * FROM users WHERE user_id = ?", (user_id,))
        user_data = cursor.fetchone()
        
        if not user_data:
            flash('用户不存在', 'danger')
            return redirect(url_for('user_management'))
        
        # 将结果转换为字典
        user = dict(zip([column[0] for column in cursor.description], user_data))
        
        # 获取用户订单信息
        cursor.execute("""
            SELECT * FROM membership_orders 
            WHERE user_id = ? 
            ORDER BY created_at DESC
        """, (user_id,))
        user_orders = []
        for row in cursor.fetchall():
            order = dict(zip([column[0] for column in cursor.description], row))
            user_orders.append(order)
        
        # 获取用户作文信息
        cursor.execute("""
            SELECT id, title, submission_time, word_count, total_score, grade
            FROM essays 
            WHERE user_id = ? 
            ORDER BY submission_time DESC
            LIMIT 10
        """, (user_id,))
        user_essays = []
        for row in cursor.fetchall():
            essay = dict(zip([column[0] for column in cursor.description], row))
            user_essays.append(essay)
        
        return render_template('admin/user_detail.html', 
                            user=user, 
                            user_orders=user_orders,
                            user_essays=user_essays)
    except Exception as e:
        logger.error(f"加载用户详情页面出错: {str(e)}")
        flash(f'加载用户详情页面出错: {str(e)}', 'danger')
        return redirect(url_for('user_management'))

# 系统配置页面
@app.route('/admin/config')
@admin_required
def website_config():
    """系统配置页面"""
    try:
        # 获取数据库连接
        conn = get_db()
        cursor = conn.cursor()
        
        # 获取所有配置项
        cursor.execute("SELECT config_name, config_value FROM website_config")
        config_rows = cursor.fetchall()
        
        # 将结果转换为字典
        config = {row[0]: row[1] for row in config_rows}
        
        # 备份信息
        backup_info = {
            'last_backup': config.get('last_backup_time', '从未备份')
        }
        
        return render_template('admin/website_config.html', 
                            config=config,
                            backup_info=backup_info,
                            active_page='config')
    except Exception as e:
        logger.error(f"加载系统配置页面出错: {str(e)}")
        flash(f'加载系统配置页面出错: {str(e)}', 'danger')
        return redirect(url_for('admin_dashboard'))

# ... 其他管理后台路由 ...

# 管理后台数据统计页面
@app.route('/admin/stats')
@admin_required
def essay_stats():
    """数据统计页面"""
    try:
        # 获取数据库连接
        conn = get_db()
        cursor = conn.cursor()
        
        # 获取作文统计数据
        cursor.execute("SELECT COUNT(*) FROM essays")
        total_essays = cursor.fetchone()[0]
        
        # 查询平均分和最高/最低分
        cursor.execute("""
            SELECT 
                ROUND(AVG(total_score), 2) as avg_score,
                MAX(total_score) as max_score,
                MIN(total_score) as min_score
            FROM essays 
            WHERE total_score IS NOT NULL
        """)
        score_stats = cursor.fetchone()
        avg_score = score_stats[0] or 0
        max_score = score_stats[1] or 0
        min_score = score_stats[2] or 0
        
        # 查询平均字数
        cursor.execute("""
            SELECT 
                ROUND(AVG(word_count), 2) as avg_words,
                MAX(word_count) as max_words
            FROM essays 
            WHERE word_count IS NOT NULL
        """)
        word_stats = cursor.fetchone()
        avg_words = word_stats[0] or 0
        max_words = word_stats[1] or 0
        
        # 按月分组的统计数据
        cursor.execute("""
            SELECT 
                strftime('%Y-%m', submission_time) as month,
                COUNT(*) as essay_count,
                ROUND(AVG(total_score), 2) as avg_score
            FROM essays
            WHERE total_score IS NOT NULL
            GROUP BY month
            ORDER BY month DESC
            LIMIT 12
        """)
        monthly_stats = [dict(zip(['month', 'count', 'avg_score'], row)) for row in cursor.fetchall()]
        
        # 构建统计数据
        stats = {
            'total_essays': total_essays,
            'avg_score': avg_score,
            'max_score': max_score,
            'min_score': min_score,
            'avg_words': avg_words,
            'max_words': max_words,
            'monthly_stats': monthly_stats,
            'score_distribution': []
        }
        
        return render_template('admin/essay_stats.html',
                            stats=stats,
                            active_page='stats')
    
    except Exception as e:
        logger.error(f"加载数据统计页面出错: {str(e)}")
        flash(f'加载数据统计页面出错: {str(e)}', 'danger')
        return redirect(url_for('admin_dashboard'))

# 更新用户基本信息
@app.route('/admin/users/<int:user_id>/update', methods=['POST'])
@admin_required
def update_user(user_id):
    """更新用户基本信息"""
    try:
        # 获取表单数据
        username = request.form.get('username')
        email = request.form.get('email')
        new_password = request.form.get('new_password')
        role = request.form.get('role')
        user_type = request.form.get('user_type')
        is_active = int(request.form.get('is_active', 1))
        
        # 获取数据库连接
        conn = get_db()
        cursor = conn.cursor()
        
        # 检查用户是否存在
        cursor.execute("SELECT 1 FROM users WHERE user_id = ?", (user_id,))
        if not cursor.fetchone():
            flash('用户不存在', 'danger')
            return redirect(url_for('user_management'))
        
        # 检查用户名是否已被其他用户使用
        cursor.execute("SELECT 1 FROM users WHERE username = ? AND user_id != ?", (username, user_id))
        if cursor.fetchone():
            flash('该用户名已被使用', 'danger')
            return redirect(url_for('user_detail', user_id=user_id))
        
        # 检查邮箱是否已被其他用户使用
        cursor.execute("SELECT 1 FROM users WHERE email = ? AND user_id != ?", (email, user_id))
        if cursor.fetchone():
            flash('该邮箱已被注册', 'danger')
            return redirect(url_for('user_detail', user_id=user_id))
        
        # 构建更新SQL
        sql = """
            UPDATE users SET 
            username = ?, 
            email = ?, 
            role = ?,
            user_type = ?,
            is_active = ?
        """
        params = [username, email, role, user_type, is_active]
        
        # 如果提供了新密码，更新密码
        if new_password:
            sql += ", password_hash = ?"
            params.append(generate_password_hash(new_password))
        
        sql += " WHERE user_id = ?"
        params.append(user_id)
        
        # 执行更新
        cursor.execute(sql, params)
        conn.commit()
        
        flash('用户信息更新成功', 'success')
        return redirect(url_for('user_detail', user_id=user_id))
    
    except Exception as e:
        logger.error(f"更新用户信息出错: {str(e)}")
        flash(f'更新用户信息出错: {str(e)}', 'danger')
        return redirect(url_for('user_detail', user_id=user_id))

# 更新用户会员权益
@app.route('/admin/users/<int:user_id>/membership', methods=['POST'])
@admin_required
def update_user_membership(user_id):
    """更新用户会员权益"""
    try:
        # 获取表单数据
        essays_remaining = request.form.get('essays_remaining', type=int)
        essays_monthly_limit = request.form.get('essays_monthly_limit', type=int)
        essays_daily_limit = request.form.get('essays_daily_limit', type=int)
        essays_daily_used = request.form.get('essays_daily_used', type=int)
        vip_status = request.form.get('vip_status', type=int)
        membership_expiry = request.form.get('membership_expiry')
        
        # 获取数据库连接
        conn = get_db()
        cursor = conn.cursor()
        
        # 检查用户是否存在
        cursor.execute("SELECT 1 FROM users WHERE user_id = ?", (user_id,))
        if not cursor.fetchone():
            flash('用户不存在', 'danger')
            return redirect(url_for('user_management'))
        
        # 执行更新
        cursor.execute("""
            UPDATE users SET 
            essays_remaining = ?,
            essays_monthly_limit = ?,
            essays_daily_limit = ?,
            essays_daily_used = ?,
            vip_status = ?,
            membership_expiry = ?
            WHERE user_id = ?
        """, (
            essays_remaining, 
            essays_monthly_limit, 
            essays_daily_limit, 
            essays_daily_used, 
            vip_status, 
            membership_expiry if membership_expiry else None, 
            user_id
        ))
        conn.commit()
        
        flash('会员权益更新成功', 'success')
        return redirect(url_for('user_detail', user_id=user_id))
    
    except Exception as e:
        logger.error(f"更新会员权益出错: {str(e)}")
        flash(f'更新会员权益出错: {str(e)}', 'danger')
        return redirect(url_for('user_detail', user_id=user_id))

# 添加新用户
@app.route('/admin/users/add', methods=['POST'])
@admin_required
def add_user():
    """添加新用户"""
    try:
        # 获取表单数据
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        user_type = request.form.get('user_type', 'free')
        role = request.form.get('role', 'user')
        
        # 获取数据库连接
        conn = get_db()
        cursor = conn.cursor()
        
        # 检查用户名是否已存在
        cursor.execute("SELECT 1 FROM users WHERE username = ?", (username,))
        if cursor.fetchone():
            flash('该用户名已被使用', 'danger')
            return redirect(url_for('user_management'))
        
        # 检查邮箱是否已存在
        cursor.execute("SELECT 1 FROM users WHERE email = ?", (email,))
        if cursor.fetchone():
            flash('该邮箱已被注册', 'danger')
            return redirect(url_for('user_management'))
        
        # 创建用户
        password_hash = generate_password_hash(password)
        now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        today = datetime.now().strftime('%Y-%m-%d')
        
        # 设置默认配置
        essays_remaining = 10
        essays_monthly_limit = 10
        essays_daily_limit = 5
        membership_expiry = None
        
        # 根据用户类型设置权益
        if user_type == 'regular':
            essays_remaining = 300
            essays_monthly_limit = 300
            essays_daily_limit = 30
            membership_expiry = (datetime.now() + timedelta(days=30)).strftime('%Y-%m-%d')
        elif user_type == 'premium':
            essays_remaining = 500
            essays_monthly_limit = 500
            essays_daily_limit = 60
            membership_expiry = (datetime.now() + timedelta(days=30)).strftime('%Y-%m-%d')
        
        cursor.execute("""
            INSERT INTO users (
                username, email, password_hash, created_at, 
                user_type, essays_remaining, essays_monthly_limit, essays_daily_limit, 
                essays_daily_used, essays_total_used, daily_reset_date, membership_expiry,
                registration_bonus_claimed, vip_status, is_active, role
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            username, email, password_hash, now,
            user_type, essays_remaining, essays_monthly_limit, essays_daily_limit,
            0, 0, today, membership_expiry,
            1, 0, 1, role
        ))
        
        conn.commit()
        
        flash('用户添加成功', 'success')
        return redirect(url_for('user_management'))
    
    except Exception as e:
        logger.error(f"添加用户出错: {str(e)}")
        flash(f'添加用户出错: {str(e)}', 'danger')
        return redirect(url_for('user_management'))

# 删除用户
@app.route('/admin/users/delete', methods=['POST'])
@admin_required
def delete_user():
    """删除用户"""
    try:
        user_id = request.form.get('user_id', type=int)
        
        # 获取数据库连接
        conn = get_db()
        cursor = conn.cursor()
        
        # 检查用户是否存在
        cursor.execute("SELECT role FROM users WHERE user_id = ?", (user_id,))
        user = cursor.fetchone()
        
        if not user:
            flash('用户不存在', 'danger')
            return redirect(url_for('user_management'))
        
        # 不允许删除管理员账号
        if user[0] == 'admin':
            flash('不能删除管理员账号', 'danger')
            return redirect(url_for('user_management'))
        
        # 删除用户相关数据
        cursor.execute("DELETE FROM essays WHERE user_id = ?", (user_id,))
        cursor.execute("DELETE FROM membership_orders WHERE user_id = ?", (user_id,))
        cursor.execute("DELETE FROM user_profiles WHERE user_id = ?", (user_id,))
        cursor.execute("DELETE FROM users WHERE user_id = ?", (user_id,))
        
        conn.commit()
        
        flash('用户已删除', 'success')
        return redirect(url_for('user_management'))
    
    except Exception as e:
        logger.error(f"删除用户出错: {str(e)}")
        flash(f'删除用户出错: {str(e)}', 'danger')
        return redirect(url_for('user_management'))

# 更新会员配置
@app.route('/admin/config/membership', methods=['POST'])
@admin_required
def update_membership_config():
    """更新会员配置"""
    try:
        # 获取表单数据
        free_essays = request.form.get('free_essays', type=int)
        free_daily_limit = request.form.get('free_daily_limit', type=int)
        regular_monthly_essays = request.form.get('regular_monthly_essays', type=int)
        regular_daily_essays = request.form.get('regular_daily_essays', type=int)
        regular_price = request.form.get('regular_price', type=float)
        premium_monthly_essays = request.form.get('premium_monthly_essays', type=int)
        premium_daily_essays = request.form.get('premium_daily_essays', type=int)
        premium_price = request.form.get('premium_price', type=float)
        registration_bonus = request.form.get('registration_bonus', type=int)
        
        # 获取数据库连接
        conn = get_db()
        cursor = conn.cursor()
        
        # 更新配置
        configs = {
            'free_essays': free_essays,
            'free_daily_limit': free_daily_limit,
            'regular_monthly_essays': regular_monthly_essays,
            'regular_daily_essays': regular_daily_essays,
            'regular_price': regular_price,
            'premium_monthly_essays': premium_monthly_essays,
            'premium_daily_essays': premium_daily_essays,
            'premium_price': premium_price,
            'registration_bonus': registration_bonus
        }
        
        for name, value in configs.items():
            cursor.execute("""
                INSERT OR REPLACE INTO website_config (config_name, config_value)
                VALUES (?, ?)
            """, (name, value))
        
        conn.commit()
        
        flash('会员配置已更新', 'success')
        return redirect(url_for('website_config'))
    
    except Exception as e:
        logger.error(f"更新会员配置出错: {str(e)}")
        flash(f'更新会员配置出错: {str(e)}', 'danger')
        return redirect(url_for('website_config'))

# 更新系统配置
@app.route('/admin/config/system', methods=['POST'])
@admin_required
def update_system_config():
    """更新系统配置"""
    try:
        # 获取表单数据
        site_title = request.form.get('site_title')
        site_description = request.form.get('site_description')
        contact_email = request.form.get('contact_email')
        maintenance_mode = request.form.get('maintenance_mode', '0')
        
        # 获取数据库连接
        conn = get_db()
        cursor = conn.cursor()
        
        # 更新配置
        configs = {
            'site_title': site_title,
            'site_description': site_description,
            'contact_email': contact_email,
            'maintenance_mode': maintenance_mode
        }
        
        for name, value in configs.items():
            cursor.execute("""
                INSERT OR REPLACE INTO website_config (config_name, config_value)
                VALUES (?, ?)
            """, (name, value))
        
        conn.commit()
        
        flash('系统配置已更新', 'success')
        return redirect(url_for('website_config'))
    
    except Exception as e:
        logger.error(f"更新系统配置出错: {str(e)}")
        flash(f'更新系统配置出错: {str(e)}', 'danger')
        return redirect(url_for('website_config'))

# 更新AI配置
@app.route('/admin/config/ai', methods=['POST'])
@admin_required
def update_ai_config():
    """更新AI配置"""
    try:
        # 获取表单数据
        ai_service_status = request.form.get('ai_service_status', '1')
        max_tokens = request.form.get('max_tokens', '4096')
        temperature = request.form.get('temperature', '0.7')
        model_name = request.form.get('model_name', 'gpt-3.5-turbo')
        
        # 获取数据库连接
        conn = get_db()
        cursor = conn.cursor()
        
        # 更新配置
        configs = {
            'ai_service_status': ai_service_status,
            'max_tokens': max_tokens,
            'temperature': temperature,
            'model_name': model_name
        }
        
        for name, value in configs.items():
            cursor.execute("""
                INSERT OR REPLACE INTO website_config (config_name, config_value)
                VALUES (?, ?)
            """, (name, value))
        
        conn.commit()
        
        flash('AI配置已更新', 'success')
        return redirect(url_for('website_config'))
    
    except Exception as e:
        logger.error(f"更新AI配置出错: {str(e)}")
        flash(f'更新AI配置出错: {str(e)}', 'danger')
        return redirect(url_for('website_config'))

# 备份数据库
@app.route('/admin/backup', methods=['POST'])
@admin_required
def backup_database():
    """备份数据库"""
    try:
        import shutil
        import os
        from datetime import datetime
        
        # 创建备份目录
        backup_dir = 'backups'
        if not os.path.exists(backup_dir):
            os.makedirs(backup_dir)
        
        # 生成备份文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_file = f"{backup_dir}/essay_correction_{timestamp}.db"
        
        # 备份数据库
        shutil.copy2('instance/essay_correction.db', backup_file)
        
        # 更新最后备份时间
        conn = get_db()
        cursor = conn.cursor()
        now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        cursor.execute("""
            INSERT OR REPLACE INTO website_config (config_name, config_value)
            VALUES ('last_backup_time', ?)
        """, (now,))
        
        conn.commit()
        
        return jsonify({
            'success': True,
            'message': f'数据库已备份到 {backup_file}'
        })
    
    except Exception as e:
        logger.error(f"备份数据库出错: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'备份失败: {str(e)}'
        })

# 清理临时文件
@app.route('/admin/clear-temp', methods=['POST'])
@admin_required
def clear_temp_files():
    """清理临时文件"""
    try:
        import os
        import glob
        
        # 清理上传目录
        upload_dir = app.config.get('UPLOAD_FOLDER', 'uploads')
        if os.path.exists(upload_dir):
            files = glob.glob(f"{upload_dir}/*")
            for file in files:
                if os.path.isfile(file):
                    os.remove(file)
        
        return jsonify({
            'success': True,
            'message': '临时文件已清理'
        })
    
    except Exception as e:
        logger.error(f"清理临时文件出错: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'清理失败: {str(e)}'
        })

# 管理后台日志查看页面
@app.route('/admin/logs')
@admin_required
def admin_log_viewer():
    """管理后台日志查看页面"""
    try:
        # 查找日志文件
        log_dir = 'logs'
        if not os.path.exists(log_dir):
            log_dir = '.'
        
        log_files_with_time = []
        for file in os.listdir(log_dir):
            if file.endswith('.log'):
                full_path = os.path.join(log_dir, file)
                log_files_with_time.append((file, os.path.getmtime(full_path)))
        
        # 按修改时间排序，最新的在前
        log_files_with_time.sort(key=lambda x: x[1], reverse=True)
        
        # 提取文件名
        log_files = [f[0] for f in log_files_with_time]
        
        selected_log = request.args.get('log')
        if not selected_log and log_files:
            selected_log = log_files[0]
        
        log_content = []
        error_count = 0
        
        if selected_log:
            log_path = os.path.join(log_dir, selected_log)
            with open(log_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
                # 只获取最后1000行
                lines = lines[-1000:]
                
                for line in lines:
                    if 'ERROR' in line or '错误' in line:
                        error_count += 1
                    log_content.append(line)
        
        return render_template('admin/log_viewer.html',
                            log_files=log_files,
                            selected_log=selected_log,
                            log_content=log_content,
                            error_count=error_count,
                            active_page='logs')
    
    except Exception as e:
        logger.error(f"加载日志查看页面出错: {str(e)}")
        flash(f'加载日志查看页面出错: {str(e)}', 'danger')
        return redirect(url_for('admin_dashboard'))

# 订单管理页面
@app.route('/admin/orders')
@admin_required
def order_management():
    """订单管理页面"""
    try:
        # 获取查询参数
        status = request.args.get('status', '')
        page = request.args.get('page', 1, type=int)
        per_page = 20  # 每页显示20条记录
        
        # 获取数据库连接
        conn = get_db()
        cursor = conn.cursor()
        
        # 构建SQL查询
        query = """
            SELECT o.*, u.username 
            FROM membership_orders o
            JOIN users u ON o.user_id = u.user_id
            WHERE 1=1
        """
        params = []
        
        if status:
            query += " AND o.payment_status = ?"
            params.append(status)
        
        # 获取总记录数
        count_query = f"SELECT COUNT(*) FROM ({query})"
        cursor.execute(count_query, params)
        total = cursor.fetchone()[0]
        
        # 计算分页信息
        total_pages = (total + per_page - 1) // per_page  # 向上取整
        offset = (page - 1) * per_page
        
        # 获取当前页的数据
        query += " ORDER BY o.created_at DESC LIMIT ? OFFSET ?"
        params.extend([per_page, offset])
        
        cursor.execute(query, params)
        rows = cursor.fetchall()
        
        # 列名
        columns = [desc[0] for desc in cursor.description]
        
        # 将结果转换为字典列表
        orders = []
        for row in rows:
            order = dict(zip(columns, row))
            orders.append(order)
        
        # 构造分页信息
        pagination = {
            'current_page': page,
            'pages': total_pages,
            'total': total
        }
        
        # 统计收入
        cursor.execute("""
            SELECT 
                SUM(amount) as total_income,
                COUNT(*) as order_count
            FROM membership_orders
            WHERE payment_status = 'completed'
        """)
        income_stats = cursor.fetchone()
        
        # 按月统计收入
        cursor.execute("""
            SELECT 
                strftime('%Y-%m', created_at) as month, 
                SUM(amount) as income,
                COUNT(*) as order_count
            FROM membership_orders
            WHERE payment_status = 'completed'
            GROUP BY month
            ORDER BY month DESC
            LIMIT 12
        """)
        monthly_income = []
        for row in cursor.fetchall():
            monthly_income.append([row[0], float(row[1]) if row[1] else 0, row[2]])
        
        # 构建统计数据
        stats = {
            'total_income': income_stats[0] or 0,
            'order_count': income_stats[1] or 0,
            'monthly_income': monthly_income
        }
        
        return render_template('admin/order_management.html',
                            orders=orders,
                            stats=stats,
                            pagination=pagination,
                            filter_status=status,
                            active_page='orders')
    
    except Exception as e:
        logger.error(f"加载订单管理页面出错: {str(e)}")
        flash(f'加载订单管理页面出错: {str(e)}', 'danger')
        return redirect(url_for('admin_dashboard'))

# 启动应用
if __name__ == '__main__':
    logger.info("【启动Flask应用】日志系统已配置完成")
    
    # 初始化Flask应用
    app = Flask(__name__)

    # 配置模板文件夹路径
    import os
    template_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), 'templates'))
    app.template_folder = template_dir
    
    # 启动Flask应用
    app.run(debug=True, host='127.0.0.1', port=5000) 